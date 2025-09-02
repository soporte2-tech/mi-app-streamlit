import streamlit as st
import time
import google.generativeai as genai
import json
import re
from pathlib import Path
from pypdf import PdfReader
import docx
import io # <-- ASEGÚRATE DE TENER ESTA LÍNEA
import os # <-- Y ESTA, PARA MANEJAR ARCHIVOS
import imgkit # <-- Y ESTA, PARA LAS IMÁGENES
from docx.shared import Inches # <-- Y ESTA, PARA EL TAMAÑO DE IMAGEN
if 'pagina_actual' not in st.session_state: st.session_state.pagina_actual = 'inicio'
if 'analisis_resultado' not in st.session_state: st.session_state.analisis_resultado = None
if 'cuestionario' not in st.session_state: st.session_state.cuestionario = None
if 'pregunta_idx' not in st.session_state: st.session_state.pregunta_idx = 0
if 'plan_de_prompts' not in st.session_state: st.session_state.plan_de_prompts = None
if 'documento_final_bytes' not in st.session_state: st.session_state.documento_final_bytes = None # Para guardar el Word

# --- 1. CONFIGURACIÓN DE LA PÁGINA Y ESTILOS (CSS) ---
st.set_page_config(layout="wide")

# CSS para un diseño limpio y profesional
css_profesional = """
    /* Fondo blanco para toda la app */
    .stApp {
        background-color: white;
    }

    /* Ocultar elementos innecesarios de Streamlit */
    .stApp > header, #MainMenu, footer {
        visibility: hidden;
    }

    /* Estilo profesional para el botón principal */
    .stButton>button {
        background-color: #0056b3;
        color: white;
        font-weight: 600;
        font-size: 1.1em;
        padding: 12px 35px;
        border: none;
        border-radius: 8px;
        box-shadow: 0 4px 10px rgba(0, 86, 179, 0.2);
        transition: all 0.3s ease;
    }
    .stButton>button:hover {
        background-color: #004494;
        transform: scale(1.05);
    }
"""
st.markdown(f"<style>{css_profesional}</style>", unsafe_allow_html=True)


# --- 2. SECCIÓN DE PROMPTS Y FUNCIONES AUXILIARES ---
PROMPT_PLANTILLA = """
Eres un analista de documentos extremadamente preciso.
Te daré el texto de una plantilla de memoria técnica y los Pliegos correspondientes.
Tu única tarea es convertirlo a un objeto JSON que contenga la estructura del indice y unas indicaciones para que la persona
que va a redactar la memoria técnica sepa todo lo necesario para poder redactar la memoria técnica con mayor puntuación.

## REGLAS ESTRICTAS:
1.  La estructura del documento debes sacarlo de la plantilla y las indicaciones mezclando esa información con la de los pliegos.
2.  El objeto JSON DEBE contener dos claves de nivel superior y solo dos: "estructura_memoria" y "matices_desarrollo".
3.  Para CADA apartado y subapartado, DEBES anteponer su numeración correspondiente (ej: "1. Título", "1.1. Subtítulo").
    ESTO ES OBLIGATORIO Y DEBE SER EN NÚMEROS NORMALES (1,2,3...) NADA DE LETRAS NI COSAS RARAS.
4.  La clave "estructura_memoria" contiene la lista de apartados y subapartados como un ÍNDICE.
    La lista "subapartados" SOLO debe contener los TÍTULOS numerados, NUNCA el texto de las instrucciones.
5.  Debes coger exactamente el mismo título del apartado o subapartado que existe en el texto de la plantilla, no lo modifiques.
    Mantenlo aunque esté en otro idioma.
6.  La clave "matices_desarrollo" desglosa CADA subapartado, asociando su título numerado con las INSTRUCCIONES completas.
    NO RESUMAS. DEBES CONTAR TODO LO QUE SEPAS DE ELLO.
    Llena estas indicaciones de mucho contexto útil para que alguien sin experiencia pueda redactar la memoria.
7.  DEBES INDICAR OBLIGATORIAMENTE LA LONGITUD DE CADA SUBAPARTADO.
    NO TE LO PUEDES INVENTAR. ESTE DATO ES CLAVE.
8.  Cada instrucción debe incluir. Si no tiene eso la instrucción no vale:
    - La longitud exacta de palabras del apartado (o aproximada según lo que se diga en los pliegos). No pongas en ningún caso
    "La longitud de este subapartado no está especificada en los documentos proporcionados", propon tú uno si no existe. Esta proposición debe
    ser coherente con el apartado que es y con lo que se valora en los pliegos.
    - Una explicación clara de lo que incluirá este apartado.
    - El objetivo de contenido para que este apartado sume a obtener la excelencia en la memoria técnica.
    - Cosas que no deben faltar en el apartado.

## MEJORAS AÑADIDAS:
- Responde SIEMPRE en formato JSON válido y bien estructurado. No incluyas texto fuera del objeto JSON.
- No inventes información: solo utiliza lo que aparezca en la plantilla o en los pliegos.
- Debes mostrar conocimiento de los pliegos, no puedes asumir que el que lee las intrucciones ya posee ese conociminento.
Debes explicar todo como si el que fuera a leer las indicaciones no supiera nada del tema y deba redactar todo el contenido.
- Mantén consistencia en la numeración (ejemplo: 1, 1.1, 1.1.1). Nunca mezcles números y letras.
- Si los pliegos mencionan tablas, gráficos o anexos obligatorios, añádelos en las indicaciones como recordatorio.
- Si hay discrepancias entre plantilla y pliego, PRIORIZA SIEMPRE lo que diga el pliego.
- Valida que cada subapartado en "estructura_memoria" tenga su correspondiente bloque en "matices_desarrollo".

## EJEMPLO DE ESTRUCTURA DE SALIDA OBLIGATORIA:
{
  "estructura_memoria": [
    {
      "apartado": "1. Análisis",
      "subapartados": ["1.1. Contexto", "1.2. DAFO"]
    }
  ],
  "matices_desarrollo": [
    {
      "apartado": "1. Análisis",
      "subapartado": "1.1. Contexto",
      "indicaciones": "El subapartado debe durar 5 páginas. Este subapartado debe describir el objeto de la contratación, que es la prestación de servicios de asesoramiento, mentoría y consultoría a personas emprendedoras autónomas en Galicia..."
    },
    {
      "apartado": "1. Análisis",
      "subapartado": "1.2. DAFO",
      "indicaciones": "El subapartado debe durar 5 páginas. Este subapartado debe conseguir mostrar ..."
    }
  ]
}
"""
PROMPT_GENERAR_LISTA_PREGUNTAS = """
Actúa como un consultor experto en licitaciones. Te proporcionaré una estructura de memoria técnica detallada con apartados, subapartados y matices (indicaciones) extraídos de unos pliegos.

Tu única tarea es convertir esta información en una lista de preguntas claras y concisas para un cliente. El objetivo es recopilar de él toda la información necesaria para poder redactar la memoria técnica y obtener la máxima puntuación.

REGLAS:
1.  Para cada subapartado relevante de la estructura, genera una o más preguntas que extraigan la información clave mencionada en las "indicaciones".
2.  Las preguntas deben ser directas y fáciles de entender para alguien que no es experto en licitaciones. Evita la jerga técnica siempre que sea posible.
3.  Tu respuesta DEBE ser un único objeto JSON válido.
4.  El JSON debe contener una única clave: "cuestionario".
5.  El valor de "cuestionario" será una lista de objetos. Cada objeto tendrá dos claves: "apartado_referencia" (el título del subapartado) y "pregunta" (la pregunta que has generado).

EJEMPLO DE ENTRADA QUE RECIBIRÁS (FRAGMENTO):
...
"matices_desarrollo": [
    {
      "apartado": "1. MEMORIA TÉCNICA",
      "subapartado": "1.1. Metodología",
      "indicaciones": "El subapartado debe durar 5 páginas. Debe describir nuestra metodología de trabajo, incluyendo las fases del proyecto, las herramientas que usaremos y cómo aseguraremos la calidad..."
    }
]
...

EJEMPLO DE SALIDA OBLIGATORIA (FRAGMENTO):
{
  "cuestionario": [
    {
      "apartado_referencia": "1.1. Metodología",
      "pregunta": "Por favor, describe paso a paso la metodología de trabajo que propondremos para este proyecto. ¿Cuáles son las fases principales?"
    },
    {
      "apartado_referencia": "1.1. Metodología",
      "pregunta": "¿Qué herramientas o software específicos utilizaremos para llevar a cabo el trabajo y por qué son las mejores para este caso?"
    }
  ]
}
"""

# --- Añádelo a tu bloque de CSS ---

css_profesional = """
    /* ... tu CSS existente ... */

    .question-box {
        border: 1px solid #e0e0e0;
        border-radius: 10px;
        padding: 20px;
        margin-top: 20px;
        margin-bottom: 20px;
        box-shadow: 0 4px 8px rgba(0, 0, 0, 0.05);
    }
"""
PROMPT_PLIEGOS = """
Eres un consultor experto en licitaciones y tu conocimiento se basa ÚNICAMENTE en los archivos que te he proporcionado.
Tu misión es analizar los Pliegos y proponer una estructura para la memoria técnica que responda a todos los requisitos y criterios de valoración.
Te daré los pliegos para hacer la memoria técnica. Revisa cuidadosamente todos los que te mando (técnicos y administrativos) para sacar la estructura obligatoria, mínima o recomendada.
Tu única tarea es convertirlo a un objeto JSON que contenga la estructura del indice y unas indicaciones para que la persona
que va a redactar la memoria técnica sepa todo lo necesario para poder redactar la memoria técnica con mayor puntuación.

## REGLAS ESTRICTAS:
1.  Tu respuesta DEBE ser un único objeto JSON válido y nada más. Sin texto introductorio ni marcadores de formato como ```json.
2.  El objeto JSON DEBE contener dos claves de nivel superior y solo dos: "estructura_memoria" y "matices_desarrollo".
3.  Para CADA apartado y subapartado, DEBES anteponer su numeración correspondiente (ej: "1. Título", "1.1. Subtítulo").
    ESTO ES OBLIGATORIO Y DEBE SER EN NÚMEROS NORMALES (1,2,3...) NADA DE LETRAS NI COSAS RARAS.
4.  La clave "estructura_memoria" contiene la lista de apartados y subapartados como un ÍNDICE.
    La lista "subapartados" SOLO debe contener los TÍTULOS numerados, NUNCA el texto de las instrucciones.
5.  Debes coger exactamente el mismo título del apartado o subapartado que existe en los Pliegos, no lo modifiques.
    Mantenlo aunque esté en otro idioma.
6.  La clave "matices_desarrollo" desglosa CADA subapartado, asociando su título numerado con las INSTRUCCIONES completas.
    NO RESUMAS. DEBES CONTAR TODO LO QUE SEPAS DE ELLO.
    Llena estas indicaciones de mucho contexto útil para que alguien sin experiencia pueda redactar la memoria.
7.  DEBES INDICAR OBLIGATORIAMENTE LA LONGITUD DE CADA SUBAPARTADO.
    NO TE LO PUEDES INVENTAR. ESTE DATO ES CLAVE.
8.  Cada instrucción debe incluir. Si no tiene eso la instrucción no vale:
    - La longitud exacta de palabras del apartado (o aproximada según lo que se diga en los Pliegos).
      No pongas en ningún caso "La longitud de este subapartado no está especificada en los documentos proporcionados";
      propone tú una si no existe. Esta proposición debe ser coherente con el apartado que es y con lo que se valora en los Pliegos.
    - Una explicación clara de lo que incluirá este apartado.
    - El objetivo de contenido para que este apartado sume a obtener la excelencia en la memoria técnica.
    - Cosas que no deben faltar en el apartado.

## MEJORAS AÑADIDAS:
- Responde SIEMPRE en formato JSON válido y bien estructurado. No incluyas texto fuera del objeto JSON.
- No inventes información: utiliza únicamente lo que aparezca en los Pliegos.
- Debes mostrar conocimiento de los Pliegos; no puedes asumir que quien lea las indicaciones ya posee ese conocimiento.
  Explica todo como si la persona que redacta no supiera nada del tema y necesitara todas las claves para escribir el contenido.
- Mantén consistencia en la numeración (ejemplo: 1, 1.1, 1.1.1). Nunca mezcles números y letras.
- Si los Pliegos mencionan tablas, gráficos o anexos obligatorios, añádelos en las indicaciones como recordatorio.
- Valida que cada subapartado en "estructura_memoria" tenga su correspondiente bloque en "matices_desarrollo".

## EJEMPLO DE ESTRUCTURA DE SALIDA OBLIGATORIA:
{
  "estructura_memoria": [
    {
      "apartado": "1. Solución Técnica",
      "subapartados": ["1.1. Metodología", "1.2. Plan de Trabajo"]
    }
  ],
  "matices_desarrollo": [
    {
      "apartado": "1. Solución Técnica",
      "subapartado": "1.1. Metodología",
      "indicaciones": "El subapartado debe durar 5 páginas. Este subapartado debe describir el objeto de la contratación, que es la prestación de servicios de asesoramiento, mentoría y consultoría a personas emprendedoras autónomas en Galicia..."
    },
    {
      "apartado": "1. Solución Técnica",
      "subapartado": "1.2. Plan de Trabajo",
      "indicaciones": "El subapartado debe durar 5 páginas. Este subapartado debe conseguir mostrar ..."
    }
  ]
}
"""

def limpiar_respuesta_json(texto_sucio: str) -> str:
    if not isinstance(texto_sucio, str): return ""
    match_bloque = re.search(r'```(?:json)?\s*(\{.*\})\s*```', texto_sucio, re.DOTALL)
    if match_bloque: return match_bloque.group(1).strip()
    match_objeto = re.search(r'\{.*\}', texto_sucio, re.DOTALL)
    if match_objeto: return match_objeto.group(0).strip()
    st.warning("No se pudo encontrar un JSON válido en la respuesta de la IA.")
    return ""

def extraer_texto_de_archivo(archivo_subido):
    texto_completo = ""
    try:
        if archivo_subido.name.endswith('.pdf'):
            reader = PdfReader(archivo_subido)
            for page in reader.pages: texto_completo += (page.extract_text() or "") + "\n"
        elif archivo_subido.name.endswith('.docx'):
            doc = docx.Document(archivo_subido)
            texto_completo = "\n".join([p.text for p in doc.paragraphs])
    except Exception as e:
        st.error(f"Error al leer el archivo '{archivo_subido.name}': {e}")
    return texto_completo

def agregar_markdown_a_word(documento, texto_markdown):
    # ... (Pega aquí tu función completa 'agregar_markdown_a_word') ...
    # ... Te la pongo aquí por si acaso ...
    patron_encabezado = re.compile(r'^(#+)\s+(.*)')
    patron_lista_numerada = re.compile(r'^\s*\d+\.\s+')
    patron_lista_viñeta = re.compile(r'^\s*[\*\-]\s+')
    def procesar_linea_con_negritas(parrafo, texto):
        partes = re.split(r'(\*\*.*?\*\*)', texto)
        for parte in partes:
            if parte.startswith('**') and parte.endswith('**'):
                parrafo.add_run(parte[2:-2]).bold = True
            elif parte:
                parrafo.add_run(parte)
    for linea in texto_markdown.split('\n'):
        linea_limpia = linea.strip()
        if not linea_limpia: continue
        match_encabezado = patron_encabezado.match(linea_limpia)
        if match_encabezado:
            nivel = len(match_encabezado.group(1))
            titulo = match_encabezado.group(2).strip()
            documento.add_heading(titulo, level=min(nivel, 4))
            continue
        if patron_lista_numerada.match(linea_limpia):
            texto_item = patron_lista_numerada.sub('', linea_limpia)
            p = documento.add_paragraph(style='List Number')
            procesar_linea_con_negritas(p, texto_item)
        elif patron_lista_viñeta.match(linea_limpia):
            texto_item = patron_lista_viñeta.sub('', linea_limpia)
            p = documento.add_paragraph(style='List Bullet')
            procesar_linea_con_negritas(p, texto_item)
        else:
            p = documento.add_paragraph()
            procesar_linea_con_negritas(p, linea_limpia)

def es_html(texto):
    patron_html = re.compile(r'<\s*html|<!DOCTYPE html>|<body|<div|<table', re.IGNORECASE)
    return bool(patron_html.search(texto))

def limpiar_respuesta_html(texto_sucio):
    match = re.search(r'```html\s*(.*?)\s*```', texto_sucio, re.DOTALL)
    if match:
        return match.group(1).strip()
    return texto_sucio.strip()

def html_a_imagen(html_content, output_filename="imagen_html.png"):
    try:
        options = {'format': 'png', 'encoding': "UTF-8", 'quiet': ''}
        imgkit.from_string(html_content, output_filename, options=options)
        if os.path.exists(output_filename):
            return output_filename
        return None
    except Exception as e:
        st.warning(f"No se pudo generar la imagen desde HTML: {e}")
        return None
        
def mostrar_resultado_analisis(data):
    """
    Muestra la estructura del análisis usando expanders para una vista más limpia.
    Los apartados principales son visibles y los subapartados están dentro de cada desplegable.
    """
    if not data:
        st.error("No se pudo generar el análisis.")
        return

    st.subheader("Estructura de Apartados Propuesta")
    st.markdown("---")

    # Bucle para crear un expander por cada apartado principal
    for seccion in data.get("estructura_memoria", []):
        apartado_principal = seccion.get('apartado', 'Sin Título')

        # Usamos st.expander para crear la sección desplegable
        with st.expander(f"**{apartado_principal}**"):
            
            subapartados = seccion.get("subapartados", [])
            
            # Si no hay subapartados, mostramos un mensaje
            if not subapartados:
                st.write("*No se encontraron subapartados para esta sección.*")
            else:
                # Si hay subapartados, los mostramos como una lista dentro del expander
                for sub in subapartados:
                    texto_limpio = sub.lstrip('- ')
                    
                    # Contamos los puntos para saber el nivel de anidación (ej: 1.1 vs 1.1.1)
                    # Esto nos permite crear una sangría visual para la jerarquía.
                    nivel = texto_limpio.count('.')
                    sangria = (nivel - 1) * 25 if nivel > 0 else 0 # 0px para nivel 1, 25px para nivel 2, etc.
                    
                    st.markdown(
                        f"<div style='margin-left: {sangria}px;'>•&nbsp; {texto_limpio}</div>", 
                        unsafe_allow_html=True
                    )

# --- 3. MANEJO DE ESTADO DE SESIÓN ---
if 'pagina_actual' not in st.session_state: st.session_state.pagina_actual = 'inicio'
if 'analisis_resultado' not in st.session_state: st.session_state.analisis_resultado = None
if 'cuestionario' not in st.session_state: st.session_state.cuestionario = None
if 'pregunta_idx' not in st.session_state: st.session_state.pregunta_idx = 0
def ir_a_fase0():
    st.session_state.pagina_actual = 'fase0'
    st.session_state.analisis_resultado = None
    st.session_state.cuestionario = None
    st.session_state.pregunta_idx = 0
def ir_al_inicio():
    st.session_state.pagina_actual = 'inicio'
    st.session_state.analisis_resultado = None
    st.session_state.cuestionario = None
    st.session_state.pregunta_idx = 0
def ir_a_fase1():
    st.session_state.pagina_actual = 'fase1'
    st.session_state.pregunta_idx = 0 # Reiniciar el índice al entrar en la fase

# --- 4. DEFINICIÓN DE LAS PÁGINAS ---
def pagina_inicio():
    _ , col2, _ = st.columns([1, 2, 1])
    with col2:
        _ , logo_col, _ = st.columns([1, 1, 1])
        with logo_col:
            try: st.image('imagen.png', width=150)
            except Exception: st.warning("⚠️ No se encontró la imagen 'imagen.png'.")
        st.markdown("<h1 style='text-align: center; color: #1E3A5F;'>AUTOMATIZACIÓN DE MEMORIAS TÉCNICAS</h1>", unsafe_allow_html=True)
        st.write("")
        st.markdown("""<p style='text-align: center; font-size: 1.1em; color: #555;'>
            Esta herramienta está diseñada para simplificar y acelerar la creación de tus memorias técnicas.<br>
            Haz clic en <b>Comenzar</b> para iniciar el proceso.</p>""", unsafe_allow_html=True)
        st.write("","")
        _ , btn_col, _ = st.columns([1, 1, 1])
        with btn_col: st.button("Comenzar", on_click=ir_a_fase0, use_container_width=True)

def pagina_fase0():
    st.title("Fase 0: Preparación y Análisis")
    if st.session_state.analisis_resultado is None:
        st.header("Sube aquí tus pliegos y plantilla para empezar el análisis")
        st.markdown("---")
        tiene_plantilla = st.radio("Plantilla:", ("Sí, voy a subir una", "No, generar solo con los pliegos"), horizontal=True, label_visibility="collapsed")
        plantilla_file = None
        if tiene_plantilla == "Sí, voy a subir una":
            plantilla_file = st.file_uploader("📂 Sube tu plantilla (DOCX o PDF)", type=['docx', 'pdf'])
        pliegos_files = st.file_uploader("📂 Sube el/los pliegos (puedes seleccionar varios)", type=['docx', 'pdf'], accept_multiple_files=True)
        st.markdown("---")

        if st.button("Analizar Documentos", use_container_width=True):
            if not pliegos_files: st.error("❌ Por favor, sube al menos un pliego.")
            elif tiene_plantilla == "Sí, voy a subir una" and not plantilla_file: st.error("❌ Has indicado que tienes plantilla, por favor, súbela.")
            else:
                with st.spinner('Realizando análisis profundo... Este proceso puede tardar varios minutos.'):
                    try:
                        api_key = st.secrets["GEMINI_API_KEY"]
                        genai.configure(api_key=api_key)
                        model = genai.GenerativeModel('gemini-1.5-flash')

                        prompt_a_usar = PROMPT_PLANTILLA if plantilla_file else PROMPT_PLIEGOS
                        contenido_ia = [prompt_a_usar]
                        
                        if plantilla_file:
                            st.info("Extrayendo texto de la plantilla...")
                            texto_plantilla = extraer_texto_de_archivo(plantilla_file)
                            if texto_plantilla: contenido_ia.append(texto_plantilla)

                        st.info("Subiendo documentos a la API...")
                        referencias_pliegos = []
                        temp_dir = Path("temp_files"); temp_dir.mkdir(exist_ok=True)
                        for pliego in pliegos_files:
                            temp_path = temp_dir / pliego.name
                            with open(temp_path, "wb") as f: f.write(pliego.getvalue())
                            referencia = genai.upload_file(path=temp_path, display_name=pliego.name)
                            referencias_pliegos.append(referencia)
                            temp_path.unlink()
                        if referencias_pliegos: contenido_ia.extend(referencias_pliegos)
                        
                        st.info("La IA está generando la estructura completa. Por favor, ten paciencia...")
                        generation_config = genai.GenerationConfig(response_mime_type="application/json", max_output_tokens=8192)
                        
                        response = model.generate_content(
                            contenido_ia,
                            generation_config=generation_config,
                            request_options={"timeout": 600}
                        )

                        if response and hasattr(response, 'text') and response.text:
                            json_limpio_str = limpiar_respuesta_json(response.text)
                            if json_limpio_str:
                                st.session_state.analisis_resultado = json.loads(json_limpio_str)
                                st.success("✅ ¡Análisis completo!")
                                time.sleep(2)
                                st.rerun()
                            else:
                                st.error("❌ La IA devolvió una respuesta, pero estaba incompleta o no era un JSON válido.")
                        else:
                            st.error(f"❌ La IA no generó una respuesta válida. Detalles: {response.prompt_feedback}")
                    except Exception as e:
                        st.error(f"Ocurrió un error inesperado durante el análisis: {e}")
    else:
        st.header("📑 Estructura Sugerida del Análisis")
        mostrar_resultado_analisis(st.session_state.analisis_resultado)
        st.markdown("---")
        col1, col2, col3 = st.columns([1,2,1])
        with col2:
            st.button("Continuar a la Fase 1", on_click=ir_a_fase1, use_container_width=True)

    st.button("Volver al Inicio", on_click=ir_al_inicio)
    
def pagina_fase1():
    st.title("Fase 1: Recopilación de Información")

    # Si aún no hemos generado el cuestionario, mostramos el botón para hacerlo
    if st.session_state.cuestionario is None:
        st.info("El siguiente paso es generar un cuestionario basado en la estructura analizada. La IA creará preguntas específicas para recopilar la información necesaria del cliente.")
        
        col1, col2, col3 = st.columns([1,2,1])
        with col2:
            if st.button("Generar Cuestionario Estratégico", use_container_width=True):
                with st.spinner("La IA está preparando las preguntas..."):
                    try:
                        api_key = st.secrets["GEMINI_API_KEY"]
                        genai.configure(api_key=api_key)
                        model = genai.GenerativeModel('gemini-1.5-flash')

                        analisis_previo = st.session_state.analisis_resultado
                        contenido_ia = [
                            PROMPT_GENERAR_LISTA_PREGUNTAS, # <-- Necesitas tener este prompt definido
                            "Aquí está la estructura y los matices generados previamente:",
                            json.dumps(analisis_previo)
                        ]

                        response = model.generate_content(
                            contenido_ia,
                            generation_config=genai.GenerationConfig(response_mime_type="application/json")
                        )

                        json_limpio_str = limpiar_respuesta_json(response.text)
                        if json_limpio_str:
                            cuestionario_data = json.loads(json_limpio_str).get("cuestionario", [])
                            for q in cuestionario_data: q['respuesta'] = ''
                            st.session_state.cuestionario = cuestionario_data
                            st.session_state.pregunta_idx = 0
                            st.rerun()
                        else:
                            st.error("La IA no pudo generar el cuestionario. Inténtalo de nuevo.")
                    except Exception as e:
                        st.error(f"Ocurrió un error: {e}")
    
    # Si ya tenemos el cuestionario, mostramos la pregunta actual
    else:
        cuestionario = st.session_state.cuestionario
        idx = st.session_state.pregunta_idx
        pregunta_actual = cuestionario[idx]
        total_preguntas = len(cuestionario)

        st.progress((idx + 1) / total_preguntas, text=f"Pregunta {idx + 1} de {total_preguntas}")

        with st.container():
            st.markdown('<div class="question-box">', unsafe_allow_html=True)
            st.subheader(f"Referente a: {pregunta_actual['apartado_referencia']}")
            st.write(pregunta_actual['pregunta'])
            respuesta = st.text_area("Respuesta del cliente:", value=pregunta_actual['respuesta'], height=200, key=f"respuesta_{idx}")
            st.markdown('</div>', unsafe_allow_html=True)

        st.session_state.cuestionario[idx]['respuesta'] = respuesta

        col1, col2, col3 = st.columns([1, 1, 1])
        with col1:
            if idx > 0:
                if st.button("⬅️ Anterior"):
                    st.session_state.pregunta_idx -= 1
                    st.rerun()
        with col3:
            if idx < total_preguntas - 1:
                if st.button("Siguiente ➡️"):
                    st.session_state.pregunta_idx += 1
                    st.rerun()
            else:
                if st.button("✅ Finalizar y Continuar a la Redacción", on_click=ir_a_fase2):
                    st.success("¡Cuestionario completado! Pasando a la fase final...")
                    time.sleep(2)
                    st.rerun()
    
    st.button("Volver al Inicio", on_click=ir_al_inicio)


def pagina_fase2():
    st.title("Fase Final: Redacción del Documento")

    if st.session_state.documento_final_bytes is None:
        st.info("La IA está lista para redactar la memoria técnica completa. Utilizará la estructura, los pliegos y tus respuestas para generar cada sección del documento.")
        st.warning("Este proceso es intensivo y puede tardar varios minutos. Por favor, no cierres esta ventana.")

        col1, col2, col3 = st.columns([1,2,1])
        with col2:
            if st.button("✍️ Iniciar Redacción Automática", use_container_width=True):
                if 'analisis_resultado' not in st.session_state or 'referencias_archivos' not in st.session_state:
                    st.error("Faltan datos de la Fase 0. Por favor, reinicia el proceso."); st.stop()

                # --- 1. GENERACIÓN DEL PLAN DE CONTENIDOS ---
                with st.spinner("Paso 1 de 2: Generando plan de contenidos..."):
                    try:
                        api_key = st.secrets["GEMINI_API_KEY"]; genai.configure(api_key=api_key)
                        model = genai.GenerativeModel('gemini-1.5-pro')
                        
                        analisis = st.session_state.analisis_resultado
                        referencias_archivos = st.session_state.referencias_archivos
                        respuestas_cliente = st.session_state.get('cuestionario', [])
                        contexto_cliente = "\n\n--- CONTEXTO ADICIONAL DEL CLIENTE ---\n"
                        for item in respuestas_cliente:
                            contexto_cliente += f"Para '{item['apartado_referencia']}', a la pregunta '{item['pregunta']}', el cliente respondió: '{item['respuesta']}'\n"

                        planes_de_accion = []
                        subapartados_a_procesar = analisis['matices_desarrollo']
                        
                        for subapartado_info in subapartados_a_procesar:
                            prompt_plan = PROMPT_DESARROLLO.format( # <-- Necesitas tu PROMPT_DESARROLLO
                                apartado_titulo=subapartado_info.get("apartado", ""),
                                subapartado_titulo=subapartado_info.get("subapartado", ""),
                                indicaciones=subapartado_info.get("indicaciones", "")
                            )
                            contenido_ia = [prompt_plan, contexto_cliente] + referencias_archivos
                            response_plan = model.generate_content(
                                contenido_ia,
                                generation_config=genai.GenerationConfig(response_mime_type="application/json"),
                                request_options={"timeout": 600}
                            )
                            json_plan = limpiar_respuesta_json(response_plan.text)
                            if json_plan:
                                planes_de_accion.extend(json.loads(json_plan).get("plan_de_prompts", []))
                        
                        st.session_state.plan_de_prompts = {"plan_de_prompts": planes_de_accion}
                    except Exception as e:
                        st.error(f"Error generando el plan de contenidos: {e}"); st.stop()

                # --- 2. ENSAMBLAJE DEL DOCUMENTO WORD ---
                plan_de_prompts = st.session_state.plan_de_prompts.get("plan_de_prompts", [])
                if not plan_de_prompts: st.error("No se pudo generar un plan de contenidos válido."); st.stop()
                
                status_placeholder = st.empty()
                progress_bar = st.progress(0)
                documento = docx.Document()
                chat_redaccion = model.start_chat()

                ultimo_apartado_escrito = ""; ultimo_subapartado_escrito = ""

                for i, tarea in enumerate(plan_de_prompts):
                    apartado = tarea.get("apartado_referencia", "")
                    subapartado = tarea.get("subapartado_referencia", "")
                    prompt_actual = tarea.get("prompt_para_asistente")
                    
                    status_placeholder.info(f"Paso 2 de 2: Redactando '{subapartado}' ({i+1}/{len(plan_de_prompts)})")
                    
                    if apartado and apartado != ultimo_apartado_escrito:
                        if ultimo_apartado_escrito: documento.add_page_break()
                        documento.add_heading(apartado, level=1)
                        ultimo_apartado_escrito = apartado; ultimo_subapartado_escrito = ""
                    if subapartado and subapartado != ultimo_subapartado_escrito:
                        documento.add_heading(subapartado, level=2)
                        ultimo_subapartado_escrito = subapartado

                    response_redaccion = chat_redaccion.send_message(prompt_actual)
                    respuesta_ia = response_redaccion.text

                    if es_html(respuesta_ia):
                        html_limpio = limpiar_respuesta_html(respuesta_ia)
                        nombre_img = f"temp_image_{i}.png"
                        image_file = html_a_imagen(html_limpio, output_filename=nombre_img)
                        if image_file:
                            documento.add_picture(image_file, width=Inches(6.0))
                            os.remove(image_file)
                    else:
                        agregar_markdown_a_word(documento, respuesta_ia)
                    
                    progress_bar.progress((i + 1) / len(plan_de_prompts))
                
                doc_io = io.BytesIO()
                documento.save(doc_io)
                st.session_state.documento_final_bytes = doc_io.getvalue()
                
                status_placeholder.success("✅ ¡Documento redactado con éxito!")
                time.sleep(2)
                st.rerun()

    else:
        st.header("🎉 ¡Tu Memoria Técnica está Lista!")
        st.info("El documento ha sido generado. Haz clic en el botón de abajo para descargarlo en formato .docx.")
        
        col1, col2, col3 = st.columns([1,2,1])
        with col2:
            st.download_button(
                label="📥 Descargar Memoria Técnica (.docx)",
                data=st.session_state.documento_final_bytes,
                file_name="Memoria_Tecnica_Generada.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )

    st.button("Empezar de Nuevo", on_click=ir_al_inicio)

# --- 5. ROUTER PRINCIPAL DE LA APLICACIÓN ---
if st.session_state.pagina_actual == 'inicio':
    pagina_inicio()
elif st.session_state.pagina_actual == 'fase0':
    pagina_fase0()
elif st.session_state.pagina_actual == 'fase1':
    pagina_fase1()
elif st.session_state.pagina_actual == 'fase2':
    pagina_fase2()
